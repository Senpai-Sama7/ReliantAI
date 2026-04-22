#!/usr/bin/env python3
"""
Enhanced Document to AI-Optimized JSON Converter v2.1.0
Converts PDF, DOCX, TXT, EPUB, and images to structured JSON for AI comprehension
Focus: Speed, accuracy, memory efficiency, AI-ready output format, and semantic cleanup
"""

import io
import json
import logging
import re
import sys
import os
import gc
import time
import mimetypes
import tempfile
import shutil
import argparse
import signal
from collections import Counter, defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from contextlib import contextmanager
from dataclasses import dataclass, asdict, field
from pathlib import Path
from typing import Dict, List, Optional, Union, Any, Tuple, Set, Iterator, Generator
from enum import Enum
import hashlib

# Configure logging with color support for console
try:
    import colorama
    colorama.init()

    class ColorFormatter(logging.Formatter):
        """Colored log formatter"""
        COLORS = {
            'DEBUG': '\033[94m',    # Blue
            'INFO': '\033[92m',     # Green
            'WARNING': '\033[93m',  # Yellow
            'ERROR': '\033[91m',    # Red
            'CRITICAL': '\033[91m\033[1m',  # Bold Red
            'RESET': '\033[0m'      # Reset
        }

        def format(self, record):
            log_message = super().format(record)
            color = self.COLORS.get(record.levelname, self.COLORS['RESET'])
            return f"{color}{log_message}{self.COLORS['RESET']}"

    # Configure colored console handler
    console_handler = logging.StreamHandler()
    console_handler.setFormatter(ColorFormatter('%(asctime)s - %(levelname)s - %(message)s'))

    # Set up logger with console handler
    logger = logging.getLogger(__name__)
    logger.setLevel(logging.INFO)
    logger.addHandler(console_handler)

except ImportError:
    # Fall back to standard logging if colorama is not available
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[logging.StreamHandler()]
    )
    logger = logging.getLogger(__name__)

BASIC_STOP_WORDS = {
    'a', 'an', 'the', 'and', 'or', 'but', 'if', 'because', 'as', 'until', 'while',
    'of', 'at', 'by', 'for', 'with', 'about', 'against', 'between', 'into', 'through',
    'during', 'before', 'after', 'above', 'below', 'to', 'from', 'up', 'down', 'in',
    'out', 'on', 'off', 'over', 'under', 'again', 'further', 'then', 'once', 'here',
    'there', 'when', 'where', 'why', 'how', 'all', 'any', 'both', 'each', 'few', 'more',
    'most', 'other', 'some', 'such', 'no', 'nor', 'not', 'only', 'own', 'same', 'so',
    'than', 'too', 'very', 'can', 'will', 'just', 'should', 'now'
}

# Version information
__version__ = "2.1.0"
logger.info(f"Document Converter v{__version__} initializing")

# Import dependencies with graceful degradation
DEPENDENCIES = {
    'core': ['fitz', 'PIL', 'pytesseract', 'docx'],
    'enhanced': ['nltk', 'langdetect', 'ebooklib', 'pdf2image'],
    'optional': ['magic', 'colorama']
}

# Dictionary to track available modules
available_modules = {}

# Core dependencies
try:
    import fitz  # PyMuPDF
    available_modules['fitz'] = True
except ImportError as e:
    logger.warning(f"PyMuPDF not available - PDF processing disabled: {e}")
    available_modules['fitz'] = False

try:
    from PIL import Image, ImageOps, ImageEnhance, ImageFilter
    available_modules['PIL'] = True
except ImportError as e:
    logger.warning(f"Pillow not available - Image processing disabled: {e}")
    available_modules['PIL'] = False

try:
    import pytesseract
    available_modules['pytesseract'] = True
except ImportError as e:
    logger.warning(f"pytesseract not available - OCR disabled: {e}")
    available_modules['pytesseract'] = False

try:
    from docx import Document as DocxDocument
    available_modules['docx'] = True
except ImportError as e:
    logger.warning(f"python-docx not available - DOCX processing disabled: {e}")
    available_modules['docx'] = False

# Enhanced dependencies
try:
    import nltk
    allow_downloads = os.environ.get('DOCUMANCER_NLTK_DOWNLOAD', '').lower() in {'1', 'true', 'yes'}
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        if allow_downloads:
            nltk.download('punkt', quiet=True)
        else:
            logger.warning("NLTK punkt not found. Set DOCUMANCER_NLTK_DOWNLOAD=1 to allow downloads.")
    try:
        nltk.data.find('corpora/stopwords')
    except LookupError:
        if allow_downloads:
            nltk.download('stopwords', quiet=True)
        else:
            logger.warning("NLTK stopwords not found. Set DOCUMANCER_NLTK_DOWNLOAD=1 to allow downloads.")
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize, sent_tokenize
    available_modules['nltk'] = True
except ImportError:
    logger.warning("NLTK not available. Topic extraction will be limited.")
    available_modules['nltk'] = False

try:
    from langdetect import detect as detect_language
    available_modules['langdetect'] = True
except ImportError:
    logger.warning("langdetect not available. Automatic language detection will be disabled.")
    available_modules['langdetect'] = False

try:
    import ebooklib
    from ebooklib import epub
    import html2text
    available_modules['ebooklib'] = True
except ImportError:
    logger.warning("ebooklib or html2text not available. EPUB processing will be disabled.")
    available_modules['ebooklib'] = False

try:
    import pdf2image
    available_modules['pdf2image'] = True
except ImportError:
    logger.warning("pdf2image not available. Alternative OCR pipeline will be disabled.")
    available_modules['pdf2image'] = False

# Optional dependencies
try:
    import magic
    available_modules['magic'] = True
except ImportError:
    logger.debug("python-magic not available. Will rely on file extensions for type detection.")
    available_modules['magic'] = False


class ContentType(Enum):
    """Content classification for AI processing"""
    HEADING = "heading"
    PARAGRAPH = "paragraph"
    CODE_BLOCK = "code_block"
    LIST = "list"
    QUOTE = "quote"
    TABLE = "table"
    METADATA = "metadata"
    FIGURE = "figure"
    EQUATION = "equation"


@dataclass
class ContentBlock:
    """Structured content block for AI consumption"""
    type: str
    content: str
    level: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary representation"""
        return {
            'type': self.type,
            'content': self.content.strip(),
            'level': self.level,
            'metadata': self.metadata or {}
        }


class ContentNormalizer:
    """Advanced content normalizer for semantic cleanup and structure improvement"""

    def __init__(self):
        self.url_cache = set()
        self.heading_patterns = self._compile_heading_patterns()
        self.language_patterns = self._compile_language_patterns()
        self.placeholder_patterns = self._compile_placeholder_patterns()

    def _compile_heading_patterns(self) -> Dict[str, re.Pattern]:
        """Compile patterns for heading normalization"""
        return {
            'chapter_variations': re.compile(r'^(Chapter|CHAPTER|Ch\.?|chapter)\s*(\d+|[IVXLCDM]+)[\s:\-]*(.*)$', re.IGNORECASE),
            'section_variations': re.compile(r'^(Section|SECTION|Sec\.?|section)\s*(\d+(?:\.\d+)*)[\s:\-]*(.*)$', re.IGNORECASE),
            'part_variations': re.compile(r'^(Part|PART|part)\s*(\d+|[IVXLCDM]+)[\s:\-]*(.*)$', re.IGNORECASE),
            'lab_step': re.compile(r'^(Lab|LAB|Step|STEP)\s*(\d+)[\s:\-]*(.*)$', re.IGNORECASE),
            'exercise': re.compile(r'^(Exercise|EXERCISE|Ex\.?)\s*(\d+)[\s:\-]*(.*)$', re.IGNORECASE),
            'duplicate_heading': re.compile(r'^(.+?)\s*\1\s*$', re.IGNORECASE),
            'placeholder_heading': re.compile(r'^[\[\(]?\s*(placeholder|todo|tbd|xxx|fixme)\s*[\]\)]?$', re.IGNORECASE)
        }

    def _compile_language_patterns(self) -> Dict[str, re.Pattern]:
        """Compile patterns for programming language detection"""
        return {
            'bash': re.compile(r'(?:#!/bin/bash|#!/bin/sh|\$\s+|sudo\s+|chmod\s+|grep\s+|awk\s+|sed\s+|ls\s+|cd\s+|mkdir\s+)', re.IGNORECASE),
            'powershell': re.compile(r'(?:PS\s*>|Get-|Set-|New-|Remove-|\$\w+\s*=|Import-Module|cmdlet)', re.IGNORECASE),
            'python': re.compile(r'(?:#!/usr/bin/python|import\s+\w+|from\s+\w+\s+import|def\s+\w+|class\s+\w+|print\()', re.IGNORECASE),
            'javascript': re.compile(r'(?:function\s+\w+|var\s+\w+|let\s+\w+|const\s+\w+|console\.log|require\()', re.IGNORECASE),
            'sql': re.compile(r'(?:SELECT\s+|FROM\s+|WHERE\s+|INSERT\s+INTO|UPDATE\s+|DELETE\s+FROM|CREATE\s+TABLE)', re.IGNORECASE),
            'yaml': re.compile(r'(?:^[\s]*[\w\-]+:\s*$|^[\s]*-\s+[\w\-]+:|version:\s*[\d\.]+)', re.MULTILINE),
            'json': re.compile(r'(?:^\s*[\{\[]|"[\w\-]+"\s*:\s*|^\s*[\}\]])', re.MULTILINE),
            'dockerfile': re.compile(r'(?:FROM\s+|RUN\s+|COPY\s+|ADD\s+|WORKDIR\s+|EXPOSE\s+)', re.IGNORECASE),
            'xml': re.compile(r'(?:<\?xml|<[\w\-]+[^>]*>|</[\w\-]+>)', re.IGNORECASE),
            'css': re.compile(r'(?:[\w\-]+\s*\{|[\w\-]+:\s*[\w\-#]+;|\.[a-zA-Z][\w\-]*\s*\{)', re.IGNORECASE)
        }

    def _compile_placeholder_patterns(self) -> List[re.Pattern]:
        """Compile patterns for placeholder detection"""
        return [
            re.compile(r'^[\[\(]?\s*(placeholder|todo|tbd|xxx|fixme|coming\s+soon)\s*[\]\)]?$', re.IGNORECASE),
            re.compile(r'^[\-_=]{10,}$'),
            re.compile(r'^\.{10,}$'),
            re.compile(r'^\s*\.\s*\.\s*\.\s*$')
        ]

    def normalize_content_blocks(self, blocks: List[ContentBlock]) -> List[ContentBlock]:
        """Apply comprehensive normalization to content blocks"""
        logger.info("Starting content normalization and semantic cleanup")

        # Phase 1: Clean individual blocks
        cleaned_blocks = []
        for block in blocks:
            normalized_block = self._normalize_individual_block(block)
            if normalized_block:
                cleaned_blocks.append(normalized_block)

        # Phase 2: Structural improvements
        structured_blocks = self._improve_structure(cleaned_blocks)

        # Phase 3: Deduplicate and merge similar content
        final_blocks = self._deduplicate_content(structured_blocks)

        logger.info(f"Normalization complete: {len(blocks)} -> {len(final_blocks)} blocks")
        return final_blocks

    def _normalize_individual_block(self, block: ContentBlock) -> Optional[ContentBlock]:
        """Normalize a single content block"""
        if not block.content.strip():
            return None

        if block.type == ContentType.HEADING.value:
            return self._normalize_heading(block)
        elif block.type == ContentType.CODE_BLOCK.value:
            return self._normalize_code_block(block)
        elif block.type == ContentType.LIST.value:
            return self._normalize_list(block)
        elif block.type == ContentType.PARAGRAPH.value:
            return self._normalize_paragraph(block)
        else:
            return block

    def _normalize_heading(self, block: ContentBlock) -> Optional[ContentBlock]:
        """Normalize heading blocks"""
        content = block.content.strip()

        if any(pattern.match(content) for pattern in self.placeholder_patterns):
            return None

        for pattern_name, pattern in self.heading_patterns.items():
            if pattern_name.endswith('_variations'):
                match = pattern.match(content)
                if match:
                    prefix = match.group(1).title()
                    number = match.group(2)
                    title = match.group(3).strip() if len(match.groups()) > 2 else ""

                    if title:
                        block.content = f"{prefix} {number}: {title}"
                    else:
                        block.content = f"{prefix} {number}"

                    block.metadata['section_type'] = prefix.lower()
                    block.metadata['section_number'] = number
                    block.metadata['normalized'] = True
                    break

        duplicate_match = self.heading_patterns['duplicate_heading'].match(content)
        if duplicate_match:
            block.content = duplicate_match.group(1).strip()
            block.metadata['had_duplicate'] = True

        if block.content.isupper() and len(block.content) > 10:
            block.content = block.content.title()
            block.metadata['case_normalized'] = True

        return block

    def _normalize_code_block(self, block: ContentBlock) -> Optional[ContentBlock]:
        """Normalize and enhance code blocks"""
        content = block.content.strip()

        if len(content) < 10 and not any(char in content for char in ['{', '}', '(', ')', ';', '=']):
            return ContentBlock(
                type=ContentType.HEADING.value,
                content=content,
                level=block.level or 3,
                metadata={**block.metadata, 'converted_from': 'code_block'}
            )

        detected_language = self._detect_programming_language(content)
        if detected_language:
            block.metadata['language'] = detected_language
            block.metadata['language_confidence'] = 'detected'

        if detected_language in ['bash', 'powershell']:
            content = self._clean_command_line_content(content)
            block.content = content

        return block

    def _normalize_list(self, block: ContentBlock) -> Optional[ContentBlock]:
        """Normalize list blocks"""
        content = block.content.strip()
        lines = content.split('\n')

        if self._is_semantic_list(lines):
            block.metadata['semantic_list'] = True
            normalized_items = []
            for line in lines:
                line = line.strip()
                if line:
                    clean_line = re.sub(r'^[\-\*\+•]\s*', '', line)
                    clean_line = re.sub(r'^\d+[\.\)]\s*', '', clean_line)
                    normalized_items.append(clean_line)

            block.content = '\n'.join(normalized_items)
            block.metadata['normalized'] = True

        return block

    def _normalize_paragraph(self, block: ContentBlock) -> Optional[ContentBlock]:
        """Normalize paragraph blocks"""
        content = block.content.strip()

        if 'urls' in block.metadata:
            cleaned_urls = []
            for url in block.metadata['urls']:
                clean_url = self._clean_url(url)
                if clean_url and clean_url not in self.url_cache:
                    cleaned_urls.append(clean_url)
                    self.url_cache.add(clean_url)
            block.metadata['urls'] = cleaned_urls

        if self._should_convert_to_list(content):
            list_items = self._extract_list_items(content)
            if list_items:
                return ContentBlock(
                    type=ContentType.LIST.value,
                    content='\n'.join(list_items),
                    level=block.level,
                    metadata={
                        **block.metadata,
                        'converted_from': 'paragraph',
                        'list_type': 'semantic',
                        'item_count': len(list_items)
                    }
                )

        block.content = content
        return block

    def _detect_programming_language(self, code: str) -> Optional[str]:
        """Detect programming language from code content"""
        if code.startswith('#!/'):
            if 'python' in code[:50]:
                return 'python'
            elif any(shell in code[:50] for shell in ['bash', 'sh']):
                return 'bash'

        for language, pattern in self.language_patterns.items():
            if pattern.search(code):
                return language

        if any(keyword in code.lower() for keyword in ['select ', 'from ', 'where ']):
            return 'sql'
        elif code.strip().startswith(('<?xml', '<html', '<!')):
            return 'xml'
        elif re.search(r'^\s*[\{\[]', code, re.MULTILINE) and '"' in code:
            return 'json'

        return 'text'

    def _clean_command_line_content(self, content: str) -> str:
        """Clean command line artifacts from code blocks"""
        lines = content.split('\n')
        cleaned_lines = []

        for line in lines:
            line = re.sub(r'^[\$#]\s*', '', line)
            line = re.sub(r'^PS\s*>\s*', '', line)
            line = re.sub(r'^\w+@\w+:\w*\$\s*', '', line)
            cleaned_lines.append(line)

        return '\n'.join(cleaned_lines)

    def _clean_url(self, url: str) -> str:
        """Clean and normalize URLs"""
        if not url:
            return ""

        url = re.sub(r'[.,;:!?\]\)]+$', '', url.strip())
        url = re.sub(r'^https?://https?://', 'https://', url)
        url = re.sub(r'^www\.', 'https://www.', url)

        if not url.startswith(('http://', 'https://', 'ftp://', 'mailto:')):
            if '.' in url and not url.startswith('/'):
                url = f'https://{url}'

        return url

    def _is_semantic_list(self, lines: List[str]) -> bool:
        """Detect if content should be formatted as a semantic list"""
        if len(lines) < 2:
            return False

        tool_indicators = ['tool', 'software', 'application', 'utility', 'command', 'package']
        if any(indicator in '\n'.join(lines).lower() for indicator in tool_indicators):
            return True

        step_indicators = ['step', 'phase', 'stage', 'procedure', 'process']
        if any(indicator in '\n'.join(lines).lower() for indicator in step_indicators):
            return True

        checklist_patterns = [
            r'^\s*[\-\*\+•]\s*\w+',
            r'^\s*\d+[\.\)]\s*\w+',
            r'^\s*[A-Z][a-z]+:\s*',
            r'^\s*\w+\s*-\s*\w+'
        ]

        pattern_matches = 0
        for line in lines:
            if any(re.match(pattern, line) for pattern in checklist_patterns):
                pattern_matches += 1

        return pattern_matches >= len(lines) * 0.6

    def _should_convert_to_list(self, content: str) -> bool:
        """Check if paragraph content should be converted to a list"""
        sentences = re.split(r'[.!?]+', content)
        if len(sentences) < 3:
            return False

        list_indicators = 0
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            if (sentence.lower().startswith(('first', 'second', 'third', 'next', 'then', 'finally')) or
                re.match(r'^\d+[\.\)]\s*', sentence) or
                sentence.startswith(('- ', '* ', '• ')) or
                ': ' in sentence and len(sentence.split(':')) == 2):
                list_indicators += 1

        return list_indicators >= len([s for s in sentences if s.strip()]) * 0.5

    def _extract_list_items(self, content: str) -> List[str]:
        """Extract list items from paragraph content"""
        items = []
        sentences = re.split(r'[.!?]+', content)

        for sentence in sentences:
            sentence = sentence.strip()
            if sentence and len(sentence) > 5:
                sentence = re.sub(r'^(first|second|third|next|then|finally)[\s,]*', '', sentence, flags=re.IGNORECASE)
                sentence = re.sub(r'^\d+[\.\)]\s*', '', sentence)
                sentence = re.sub(r'^[\-\*\+•]\s*', '', sentence)

                if sentence:
                    items.append(sentence.strip())

        return items

    def _improve_structure(self, blocks: List[ContentBlock]) -> List[ContentBlock]:
        """Improve overall document structure"""
        improved_blocks = []
        current_context = {'chapter': None, 'section': None, 'subsection': None}

        for block in blocks:
            if block.type == ContentType.HEADING.value:
                self._update_context(block, current_context)

            block.metadata['structural_context'] = current_context.copy()

            if (block.type == ContentType.PARAGRAPH.value and
                improved_blocks and
                improved_blocks[-1].type == ContentType.PARAGRAPH.value and
                len(block.content) < 100 and len(improved_blocks[-1].content) < 100):

                improved_blocks[-1].content += " " + block.content
                if 'urls' in block.metadata:
                    prev_urls = improved_blocks[-1].metadata.get('urls', [])
                    improved_blocks[-1].metadata['urls'] = prev_urls + block.metadata['urls']
                continue

            improved_blocks.append(block)

        return improved_blocks

    def _update_context(self, block: ContentBlock, context: Dict[str, str]):
        """Update structural context based on heading"""
        level = block.level
        content = block.content

        if level == 1 or 'chapter' in block.metadata.get('section_type', ''):
            context['chapter'] = content
            context['section'] = None
            context['subsection'] = None
        elif level == 2 or 'section' in block.metadata.get('section_type', ''):
            context['section'] = content
            context['subsection'] = None
        elif level >= 3:
            context['subsection'] = content

    def _deduplicate_content(self, blocks: List[ContentBlock]) -> List[ContentBlock]:
        """Remove duplicate content blocks"""
        seen_content = set()
        unique_blocks = []

        for block in blocks:
            content_hash = hashlib.md5(block.content.encode('utf-8')).hexdigest()

            if block.type == ContentType.HEADING.value:
                normalized_content = re.sub(r'\s+', ' ', block.content.lower().strip())
                semantic_hash = hashlib.md5(normalized_content.encode('utf-8')).hexdigest()

                if semantic_hash in seen_content:
                    logger.debug(f"Removing duplicate heading: {block.content}")
                    continue
                seen_content.add(semantic_hash)
            else:
                if content_hash in seen_content:
                    logger.debug(f"Removing duplicate content block")
                    continue
                seen_content.add(content_hash)

            unique_blocks.append(block)

        return unique_blocks


class ProgressTracker:
    """Tracks and reports progress of long-running operations"""

    def __init__(self, total_items: int, description: str = "Processing"):
        self.total = total_items
        self.current = 0
        self.start_time = time.time()
        self.description = description
        self.last_report_time = 0
        self.report_interval = 1.0

    def update(self, items_completed: int = 1) -> None:
        """Update progress and print status if interval elapsed"""
        self.current += items_completed
        current_time = time.time()

        if (self.current == 1 or
            self.current == self.total or
            current_time - self.last_report_time > self.report_interval):

            percentage = min(100, int((self.current / self.total) * 100))
            elapsed = current_time - self.start_time

            if self.current > 0 and elapsed > 0:
                items_per_sec = self.current / elapsed
                eta = (self.total - self.current) / items_per_sec if items_per_sec > 0 else 0
                eta_str = f", ETA: {int(eta)}s" if self.current < self.total else ""

                # Print progress in a format that can be parsed by the Electron frontend
                print(f"Processing: {percentage}% ({self.current}/{self.total}{eta_str})", flush=True)
                logger.info(f"{self.description}: {percentage}% ({self.current}/{self.total}{eta_str})")
            else:
                print(f"Processing: {percentage}% ({self.current}/{self.total})", flush=True)
                logger.info(f"{self.description}: {percentage}% ({self.current}/{self.total})")

            self.last_report_time = current_time


@contextmanager
def memory_management():
    """Context manager for controlled memory cleanup"""
    try:
        yield
    finally:
        gc.collect()


class DocumentParser:
    """Advanced parser for AI-readable structured content with enhanced normalization"""

    def __init__(self, language: str = 'auto'):
        self.patterns = self._compile_patterns()
        self.code_indicators = {
            'prefixes': ['$', '#', '>>>', '>', 'C:\\', '~/', 'PS>', '\u03BB '],
            'extensions': ['.py', '.js', '.html', '.css', '.sql', '.sh', '.bat', '.c', '.cpp', '.java', '.rb', '.php'],
            'keywords': [
                'def ', 'function ', 'class ', 'import ', 'from ', 'SELECT ', 'FROM ', 'WHERE ', 'var ',
                'const ', 'let ', 'public ', 'private ', 'static ', 'void ', 'int ', 'string ', 'bool ',
                'package ', '@Override', '#include', '#define', 'module', 'namespace'
            ]
        }

        self.language = language
        self.normalizer = ContentNormalizer()
        self._initialize_nlp_resources()

    def _initialize_nlp_resources(self):
        """Initialize NLP resources for text analysis"""
        self.stop_words = set(BASIC_STOP_WORDS)

        if available_modules.get('nltk', False):
            try:
                self.stop_words = set(stopwords.words('english'))
                logger.debug("NLTK stopwords loaded successfully")
            except Exception as e:
                logger.warning(f"Failed to load NLTK stopwords: {e}")
        else:
            logger.debug("Using basic stopwords set (NLTK not available)")
            self.stop_words = set(BASIC_STOP_WORDS)

    def _compile_patterns(self) -> Dict[str, re.Pattern]:
        """Compile regex patterns for content detection"""
        return {
            'heading_hash': re.compile(r'^(#{1,6})\s+(.+)$'),
            'heading_underline': re.compile(r'^(.+)\n([=\-])\2{2,}$', re.MULTILINE),
            'heading_numbered': re.compile(r'^(\d+(?:\.\d+)*)\s+(.+)$'),
            'heading_chapter': re.compile(r'^(Chapter|Section|Part|CHAPTER|SECTION|PART|Ch\.?|Sec\.?)\s*(\d+|[IVXLCDM]+)[\s:\-]*(.*)$', re.IGNORECASE),
            'code_fence': re.compile(r'^```(\w+)?\n(.*?)^```$', re.MULTILINE | re.DOTALL),
            'code_indent': re.compile(r'^(    .+)$', re.MULTILINE),
            'code_line': re.compile(r'^[\$#>]\s*(.+)$'),
            'bullet_list': re.compile(r'^[\s]*[-\*\+\u2022\u25E6\u25D8\u25CB\u25CF]\s+(.+)$'),
            'numbered_list': re.compile(r'^[\s]*(\d+|[a-z]|[A-Z]|[ivxlcdm]+|[IVXLCDM]+)[.)\]]\s+(.+)$'),
            'blockquote': re.compile(r'^>\s*(.+)$'),
            'table_row': re.compile(r'^\|(.+)\|$'),
            'table_separator': re.compile(r'^[\|\+][-\+\|]+[\|\+]$'),
            'figure_caption': re.compile(r'^(Figure|Fig\.?|Table|Tbl\.?)\s*(\d+(?:\.\d+)*)[\.:]\s*(.+)$', re.IGNORECASE),
            'equation': re.compile(r'^\s*\$\$(.*?)\$\$\s*$', re.DOTALL),
            'url': re.compile(r'https?://[^\s<>"{}|\\^`[\]]+'),
            'email': re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'),
            'file_path': re.compile(r'[/\\]?[\w\-./\\]+\.\w{1,6}'),
            'date': re.compile(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}\b'),
        }

    def parse(self, text: str, source_type: str = None) -> Dict[str, Any]:
        """Parse text into AI-optimized JSON structure with enhanced normalization"""
        word_count = len(text.split())
        logger.info(f"Parsing {word_count} words of content")

        detected_lang = 'en'
        if self.language == 'auto' and available_modules.get('langdetect', False):
            try:
                sample_size = min(5000, len(text))
                sample = text[:sample_size]
                detected_lang = detect_language(sample)
                logger.info(f"Detected document language: {detected_lang}")

                if available_modules.get('nltk', False):
                    try:
                        self.stop_words = set(stopwords.words(detected_lang))
                    except Exception:
                        logger.debug(f"No NLTK stopwords available for {detected_lang}, using English")
            except Exception as e:
                logger.warning(f"Language detection failed: {e}")

        progress = ProgressTracker(5, "Text processing")

        with memory_management():
            text = self._normalize_text(text)
            progress.update()

        with memory_management():
            sections = self._split_into_sections(text)
            progress.update()

        content_blocks = []
        total_sections = len(sections)
        section_progress = ProgressTracker(total_sections, "Processing sections")

        for section in sections:
            with memory_management():
                blocks = self._process_section(section, source_type)
                content_blocks.extend(blocks)
                section_progress.update()

        progress.update()

        with memory_management():
            content_blocks = self.normalizer.normalize_content_blocks(content_blocks)
            progress.update()

        with memory_management():
            metadata = self._extract_document_metadata(content_blocks)
            metadata['language'] = detected_lang
            all_text = " ".join([block.content for block in content_blocks
                              if block.type in [ContentType.PARAGRAPH.value, ContentType.HEADING.value]])
            key_topics = self._extract_key_topics(all_text)
            progress.update()

        result = {
            'document_type': 'structured_text',
            'metadata': metadata,
            'content_blocks': [block.to_dict() for block in content_blocks],
            'summary': self._generate_summary(content_blocks, key_topics),
            'version': __version__,
            'normalization_applied': True
        }

        logger.info(f"Parsed into {len(content_blocks)} content blocks with {len(key_topics)} key topics")
        return result

    def _normalize_text(self, text: str) -> str:
        """Normalize text for consistent processing"""
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        # Normalize smart quotes to regular quotes
        text = re.sub(r'[\u201C\u201D\u201E\u201F\u2033\u2036]', '"', text)
        text = re.sub(r'[\u2018\u2019\u201A\u201B\u2032\u2035]', "'", text)
        # Normalize dashes and ellipsis
        text = re.sub(r'[\u2014\u2013\u2012]', '-', text)
        text = re.sub(r'\u2026', '...', text)
        text = re.sub(r'(\w+)-\n(\w+)', r'\1\2', text)
        text = re.sub(r'(\w)\.(\n)(\w)', r'\1. \3', text)
        text = re.sub(r'[\u00A0\u2000-\u200B\u202F\u205F\u3000]', ' ', text)
        text = re.sub(r'[\u2022\u25E6\u25D8\u25CB\u25CF]', '\u2022', text)

        return text.strip()

    def _split_into_sections(self, text: str) -> List[str]:
        """Split text into logical sections"""
        section_pattern = re.compile(r'\n\n+')
        sections = section_pattern.split(text)
        return [s.strip() for s in sections if s.strip()]

    def _process_section(self, section: str, source_type: str = None) -> List[ContentBlock]:
        """Process a section and return content blocks"""
        blocks = []

        if self._is_heading(section):
            level, content = self._extract_heading(section)
            blocks.append(ContentBlock(
                type=ContentType.HEADING.value,
                content=content,
                level=level,
                metadata={'source_type': source_type} if source_type else {}
            ))
        elif self._is_code_block(section):
            language = self._detect_code_language(section)
            blocks.append(ContentBlock(
                type=ContentType.CODE_BLOCK.value,
                content=section,
                metadata={'language': language, 'source_type': source_type} if source_type else {'language': language}
            ))
        elif self._is_list(section):
            blocks.append(ContentBlock(
                type=ContentType.LIST.value,
                content=section,
                metadata={'source_type': source_type} if source_type else {}
            ))
        elif self._is_quote(section):
            blocks.append(ContentBlock(
                type=ContentType.QUOTE.value,
                content=section.lstrip('> '),
                metadata={'source_type': source_type} if source_type else {}
            ))
        elif self._is_table(section):
            blocks.append(ContentBlock(
                type=ContentType.TABLE.value,
                content=section,
                metadata={'source_type': source_type} if source_type else {}
            ))
        else:
            urls = self.patterns['url'].findall(section)
            metadata = {'source_type': source_type} if source_type else {}
            if urls:
                metadata['urls'] = urls

            blocks.append(ContentBlock(
                type=ContentType.PARAGRAPH.value,
                content=section,
                metadata=metadata
            ))

        return blocks

    def _is_heading(self, text: str) -> bool:
        """Check if text is a heading"""
        text = text.strip()

        if self.patterns['heading_hash'].match(text):
            return True
        if self.patterns['heading_chapter'].match(text):
            return True
        if self.patterns['heading_numbered'].match(text):
            return True
        if len(text) < 100 and text.isupper():
            return True
        if len(text) < 80 and not text.endswith(('.', '?', '!', ':')):
            return True

        return False

    def _extract_heading(self, text: str) -> Tuple[int, str]:
        """Extract heading level and content"""
        text = text.strip()

        hash_match = self.patterns['heading_hash'].match(text)
        if hash_match:
            level = len(hash_match.group(1))
            content = hash_match.group(2)
            return level, content

        chapter_match = self.patterns['heading_chapter'].match(text)
        if chapter_match:
            return 1, text

        numbered_match = self.patterns['heading_numbered'].match(text)
        if numbered_match:
            number = numbered_match.group(1)
            level = number.count('.') + 1
            return min(level, 6), text

        return 2, text

    def _is_code_block(self, text: str) -> bool:
        """Check if text is a code block"""
        text = text.strip()

        if text.startswith('```') and text.endswith('```'):
            return True

        lines = text.split('\n')
        code_indicators = 0

        for line in lines:
            if line.startswith('    ') or line.startswith('\t'):
                code_indicators += 1
            elif any(line.strip().startswith(prefix) for prefix in self.code_indicators['prefixes']):
                code_indicators += 1
            elif any(keyword in line for keyword in self.code_indicators['keywords']):
                code_indicators += 1

        return code_indicators >= len(lines) * 0.5

    def _detect_code_language(self, code: str) -> str:
        """Detect the programming language of code"""
        code = code.strip()

        if code.startswith('```'):
            first_line = code.split('\n')[0]
            lang = first_line[3:].strip()
            if lang:
                return lang

        return self.normalizer._detect_programming_language(code)

    def _is_list(self, text: str) -> bool:
        """Check if text is a list"""
        lines = text.strip().split('\n')
        if len(lines) < 2:
            return False

        list_lines = 0
        for line in lines:
            if self.patterns['bullet_list'].match(line) or self.patterns['numbered_list'].match(line):
                list_lines += 1

        return list_lines >= len(lines) * 0.5

    def _is_quote(self, text: str) -> bool:
        """Check if text is a blockquote"""
        lines = text.strip().split('\n')
        quote_lines = sum(1 for line in lines if line.startswith('>'))
        return quote_lines >= len(lines) * 0.5

    def _is_table(self, text: str) -> bool:
        """Check if text is a table"""
        lines = text.strip().split('\n')
        if len(lines) < 2:
            return False

        table_lines = 0
        for line in lines:
            if self.patterns['table_row'].match(line) or self.patterns['table_separator'].match(line):
                table_lines += 1

        return table_lines >= len(lines) * 0.5

    def _extract_document_metadata(self, blocks: List[ContentBlock]) -> Dict[str, Any]:
        """Extract document metadata from content blocks"""
        metadata = {
            'title': None,
            'word_count': 0,
            'block_count': len(blocks),
            'processed_date': time.strftime('%Y-%m-%dT%H:%M:%SZ'),
            'content_types': {}
        }

        for block in blocks:
            metadata['word_count'] += len(block.content.split())

            block_type = block.type
            if block_type not in metadata['content_types']:
                metadata['content_types'][block_type] = 0
            metadata['content_types'][block_type] += 1

            if block.type == ContentType.HEADING.value and block.level == 1 and not metadata['title']:
                metadata['title'] = block.content

        return metadata

    def _extract_key_topics(self, text: str) -> List[str]:
        """Extract key topics from text"""
        if not text:
            return []

        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())

        filtered_words = [w for w in words if w not in self.stop_words]

        word_freq = Counter(filtered_words)

        top_words = word_freq.most_common(20)

        return [word for word, count in top_words if count > 1]

    def _generate_summary(self, blocks: List[ContentBlock], key_topics: List[str]) -> str:
        """Generate a summary of the document"""
        headings = [b.content for b in blocks if b.type == ContentType.HEADING.value][:5]

        first_paragraph = None
        for block in blocks:
            if block.type == ContentType.PARAGRAPH.value and len(block.content) > 50:
                first_paragraph = block.content[:300]
                if len(block.content) > 300:
                    first_paragraph += "..."
                break

        summary_parts = []

        if headings:
            summary_parts.append(f"Document contains {len(headings)} main sections")

        if key_topics:
            summary_parts.append(f"Key topics: {', '.join(key_topics[:5])}")

        if first_paragraph:
            summary_parts.append(f"Beginning: {first_paragraph}")

        return " | ".join(summary_parts) if summary_parts else "Document processed successfully"


class DocumentConverter:
    """Main document converter class that handles multiple file formats"""

    SUPPORTED_FORMATS = {
        'pdf': ['.pdf'],
        'docx': ['.docx', '.doc'],
        'text': ['.txt', '.text', '.md', '.markdown', '.rst'],
        'epub': ['.epub'],
        'image': ['.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.gif']
    }

    def __init__(self, output_dir: str = None, language: str = 'auto'):
        self.output_dir = Path(output_dir) if output_dir else Path.cwd() / 'output'
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.parser = DocumentParser(language=language)
        self.language = language

    def get_file_type(self, file_path: Path) -> Optional[str]:
        """Determine file type from extension or magic bytes"""
        ext = file_path.suffix.lower()

        for file_type, extensions in self.SUPPORTED_FORMATS.items():
            if ext in extensions:
                return file_type

        if available_modules.get('magic', False):
            try:
                mime = magic.from_file(str(file_path), mime=True)
                if 'pdf' in mime:
                    return 'pdf'
                elif 'word' in mime or 'document' in mime:
                    return 'docx'
                elif 'epub' in mime:
                    return 'epub'
                elif 'image' in mime:
                    return 'image'
                elif 'text' in mime:
                    return 'text'
            except Exception as e:
                logger.warning(f"Magic detection failed: {e}")

        return None

    def convert_file(self, file_path: Union[str, Path]) -> Optional[Path]:
        """Convert a single file to JSON"""
        file_path = Path(file_path)

        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None

        file_type = self.get_file_type(file_path)
        if not file_type:
            logger.error(f"Unsupported file format: {file_path.suffix}")
            return None

        logger.info(f"Converting {file_path.name} ({file_type})")

        try:
            text = self._extract_text(file_path, file_type)

            if not text or not text.strip():
                logger.warning(f"No text extracted from {file_path.name}")
                return None

            result = self.parser.parse(text, source_type=file_type)

            result['metadata']['source_file'] = file_path.name
            result['metadata']['source_format'] = file_type
            result['metadata']['source_path'] = str(file_path.absolute())

            output_path = self.output_dir / f"{file_path.stem}.json"

            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(result, f, indent=2, ensure_ascii=False)

            logger.info(f"Converted: {file_path.name} -> {output_path.name}")
            print(f"Converted: {file_path.name} -> {output_path}", flush=True)

            return output_path

        except Exception as e:
            logger.error(f"Error converting {file_path.name}: {e}")
            return None

    def _extract_text(self, file_path: Path, file_type: str) -> str:
        """Extract text from file based on type"""
        extractors = {
            'pdf': self._extract_pdf,
            'docx': self._extract_docx,
            'text': self._extract_text_file,
            'epub': self._extract_epub,
            'image': self._extract_image_ocr
        }

        extractor = extractors.get(file_type)
        if not extractor:
            raise ValueError(f"No extractor for file type: {file_type}")

        return extractor(file_path)

    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        if not available_modules.get('fitz', False):
            raise ImportError("PyMuPDF not available for PDF processing")

        text_parts = []

        with fitz.open(str(file_path)) as doc:
            total_pages = len(doc)
            progress = ProgressTracker(total_pages, "Extracting PDF pages")

            for page_num, page in enumerate(doc):
                page_text = page.get_text()

                if not page_text.strip() and available_modules.get('pytesseract', False):
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    page_text = pytesseract.image_to_string(img)

                text_parts.append(page_text)
                progress.update()

        return '\n\n'.join(text_parts)

    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        if not available_modules.get('docx', False):
            raise ImportError("python-docx not available for DOCX processing")

        doc = DocxDocument(str(file_path))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                style_name = para.style.name if para.style else ""

                if 'Heading' in style_name:
                    level = 1
                    try:
                        level = int(style_name.split()[-1])
                    except (ValueError, IndexError):
                        pass
                    text_parts.append(f"{'#' * level} {para.text}")
                else:
                    text_parts.append(para.text)

        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells)
                table_text.append(f"| {row_text} |")
            text_parts.append('\n'.join(table_text))

        return '\n\n'.join(text_parts)

    def _extract_text_file(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()

    def _extract_epub(self, file_path: Path) -> str:
        """Extract text from EPUB file"""
        if not available_modules.get('ebooklib', False):
            raise ImportError("ebooklib not available for EPUB processing")

        book = epub.read_epub(str(file_path))
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = True
        h.ignore_emphasis = False

        text_parts = []

        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                content = item.get_content().decode('utf-8', errors='replace')
                text = h.handle(content)
                text_parts.append(text)

        return '\n\n'.join(text_parts)

    def _extract_image_ocr(self, file_path: Path) -> str:
        """Extract text from image using OCR"""
        if not available_modules.get('PIL', False):
            raise ImportError("Pillow not available for image processing")
        if not available_modules.get('pytesseract', False):
            raise ImportError("pytesseract not available for OCR")

        img = Image.open(file_path)

        if img.mode != 'RGB':
            img = img.convert('RGB')

        img = ImageOps.autocontrast(img)

        enhancer = ImageEnhance.Sharpness(img)
        img = enhancer.enhance(2.0)

        text = pytesseract.image_to_string(img, lang='eng')

        return text

    def convert_batch(self, file_paths: List[Union[str, Path]], max_workers: int = 4) -> Dict[str, Any]:
        """Convert multiple files in parallel"""
        results = {
            'successful': [],
            'failed': [],
            'total': len(file_paths)
        }

        progress = ProgressTracker(len(file_paths), "Converting files")

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_path = {executor.submit(self.convert_file, path): path for path in file_paths}

            for future in as_completed(future_to_path):
                path = future_to_path[future]
                try:
                    result = future.result()
                    if result:
                        results['successful'].append({
                            'input': str(path),
                            'output': str(result)
                        })
                    else:
                        results['failed'].append({
                            'input': str(path),
                            'error': 'Conversion returned None'
                        })
                except Exception as e:
                    results['failed'].append({
                        'input': str(path),
                        'error': str(e)
                    })

                progress.update()

        return results


def parse_arguments():
    """Parse command line arguments"""
    parser = argparse.ArgumentParser(
        description='DocuMancer - Advanced Document to AI-Optimized JSON Converter',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  %(prog)s document.pdf
  %(prog)s --output-dir ./output *.docx *.pdf
  %(prog)s --format json --language en document.epub
        '''
    )

    parser.add_argument(
        'files',
        nargs='+',
        help='Files to convert (supports glob patterns)'
    )

    parser.add_argument(
        '-o', '--output-dir',
        default=None,
        help='Output directory for converted files (default: ./output)'
    )

    parser.add_argument(
        '-f', '--format',
        choices=['json'],
        default='json',
        help='Output format (default: json)'
    )

    parser.add_argument(
        '-l', '--language',
        default='auto',
        help='Document language for processing (default: auto-detect)'
    )

    parser.add_argument(
        '-w', '--workers',
        type=int,
        default=4,
        help='Number of parallel workers for batch processing (default: 4)'
    )

    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Enable verbose logging'
    )

    parser.add_argument(
        '--version',
        action='version',
        version=f'DocuMancer v{__version__}'
    )

    return parser.parse_args()


def main():
    """Main entry point for the converter"""
    args = parse_arguments()

    if args.verbose:
        logger.setLevel(logging.DEBUG)

    logger.info(f"DocuMancer v{__version__} starting")

    output_dir = args.output_dir or os.path.join(os.path.expanduser('~'), 'DocuMancer_Output')

    converter = DocumentConverter(
        output_dir=output_dir,
        language=args.language
    )

    file_paths = []
    for pattern in args.files:
        path = Path(pattern)
        if path.is_dir():
            logger.warning(f"Skipping directory: {path}. To process files inside, use a glob pattern like '{path}/*'")
            continue
        if path.exists():
            file_paths.append(path)
        else:
            import glob
            matched = glob.glob(pattern)
            # Filter out directories from glob results
            file_paths.extend(Path(m) for m in matched if not Path(m).is_dir())

    if not file_paths:
        logger.error("No valid files found to convert")
        sys.exit(1)

    logger.info(f"Found {len(file_paths)} file(s) to convert")

    if len(file_paths) == 1:
        result = converter.convert_file(file_paths[0])
        if result:
            logger.info(f"Successfully converted to: {result}")
            sys.exit(0)
        else:
            logger.error("Conversion failed")
            sys.exit(1)
    else:
        results = converter.convert_batch(file_paths, max_workers=args.workers)

        logger.info(f"Conversion complete: {len(results['successful'])}/{results['total']} successful")

        if results['failed']:
            logger.warning(f"Failed conversions: {len(results['failed'])}")
            for failure in results['failed']:
                logger.warning(f"  - {failure['input']}: {failure['error']}")

        sys.exit(0 if not results['failed'] else 1)


if __name__ == '__main__':
    main()
